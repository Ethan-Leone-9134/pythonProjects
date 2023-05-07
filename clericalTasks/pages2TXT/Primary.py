
import os
import pydoc
import time
import docx
from docx import Document
import convertapi
import signal
import PyPDF2
import subprocess

def getTextWord(fileName):
    # Function getTextWord
    # Pulls the text from a word document
    # Inputs  : fileName - filePath
    # Outputs : combined - text from document
    #           valid - is the input viable for change


    if os.path.exists(fileName):
        doc = docx.Document(fileName)  # Get file information
        fullTxt = []  # Initialize text storing variable
        valid = 1  # Initialize varied font logical
        # print("HI DAD")
        for imageI in doc.inline_shapes:
            # print("SHUT UP MEG")
            if imageI.height > 0:
                return "", 0


        for temp in doc.paragraphs:  # For each paragraph
            fullTxt.append(temp.text)  # Add text to variable
            for temp2 in temp.runs:  # For each paragraph's font data
                # print(temp.text)
                if temp2.font.bold is True:
                    valid = valid + 2
                if temp2.font.italic is True:
                    print(filePath + ": BAD - ITALIC")
                    return "", 0
                if temp2.font.underline is True:
                    print(filePath + ": BAD - UNDERLINE")
                    return "", 0
                # print(temp.alignment)
                if "CENTER" in str(temp.alignment):
                    print(filePath + ": BAD - ALIGNMENT")
                    return "", 0
        if valid >= 11:  # This runs if there is a t/b bold thing
            fullTxt = []
            for temp in doc.paragraphs:  # For each paragraph
                for temp2 in temp.runs:  # For each paragraph's font data
                    if temp2.font.bold is True:
                        fullTxt.append("T: " + str(temp.text))
                        break
                else:
                    fullTxt.append("B: " + str(temp.text))  # Add text to variable
        elif valid > 1:  # Kill because bold
            print(filePath + ": BAD - BOLD")
            os.startfile(filePath)
            goInp = input("Would you like to modify or leave as is? (M or L)")
            if (goInp == "L") or (goInp == "l"):
                return "", 0
            # wordapp.filePath.Close()
            # os.kill(os.getpid(filePath), signal.SIGSTOP)


        combined = '\n'.join(fullTxt)  # Merge the lines into a single document
        return combined, 1
    return "", 0


def getTextPages(fileName):
    txt = PyPDF2.PdfFileReader(fileName)
    return txt


def word2txtDoc(fileName):
    # Function word2txtDoc
    # Converts a word document into the desired text file format
    # Inputs  : fileName - File path for the word doc
    # Outputs : NONE

    [oldText, valid] = getTextWord(fileName)
    # print(valid)
    if valid == 1:
        charCount = (len(oldText))
        if charCount <= 2:
            os.remove(fileName)  # Remove the empty file
            return  # End loop
        elif charCount <= 300:
            newPath = fileName.replace(".docx", ", prompt.txt")  # File Path for new short files
            try:
                newDoc = open(newPath, "w")
                newDoc.write(oldText)
                newDoc.close()
            except:
                newDoc.close()
                os.remove(newPath)
                print(fileName + " : BAD")
            else:
                os.remove(fileName)
                os.remove(newPath)
                addPrompt(newPath, oldText)
                print(fileName + " : GOOD")
        else:
            newPath = fileName.replace(".docx", ".txt")  # File Path for new files
            try:
                newDoc = open(newPath, "w")
                newDoc.write(oldText)
            except:
                print(fileName + " : BAD")
            else:
                os.remove(fileName)
                print(fileName + " : GOOD")


def getDocx(folderPath):
    # Function getDocx
    # Calls all docx files in a folder
    # Inputs  : folderPath - Pathway for the folder
    # Outputs : mainList - Array containing the desired folders

    wholeList = (os.listdir(folderPath))  # Call entire folder contents
    mainList = []  # Initialize list variable
    for doc in wholeList:  # Repeat with each item
        if doc[-5:] == ".docx":  # If item is .docx
            mainList.append(doc)  # Add good item to list
    return mainList


def addPrompt(filePath, txt):

    promptFile = []  # New file text
    chopped = filePath.split(chr(92))
    folderPathway = chr(92).join(chopped[0:-1])  # Folder pathway
    fileName = chopped[-1]  # Name of file
    newName = fileName.replace(", prompt.txt", "")  # Erase file type
    promptFile.append(newName + ":")  # Add og title

    if txt[0:1] == "\n":  # If line starts with a line break
        txt = txt[1:len(txt)]  # Remove line break
    promptFile.append(txt)
    promptFile.append("")  # Add line after each iteration
    promptFile.append("")  # Add line after each iteration

    newDoc = open(folderPathway + chr(92) + "Prompts.txt", "a")
    # print(promptFile)
    fileTxt = '\n'.join(promptFile)
    # print(fileTxt)
    newDoc.write(fileTxt)


def initializePrompt(folderPathway):  # Create/open the found prompt document
    proDocPath = folderPathway + chr(92) + "Prompts.docx"
    # txtDoc = open(folderPathway + chr(92) + "Prompts.txt", "w")  # Makes sure there is a prompt text doc
    # txtDoc.close()
    txtDoc = open(folderPathway + chr(92) + "Prompts.txt", "a")  # Makes sure there is a prompt text doc
    if os.path.exists(proDocPath):
        [txt, valid] = getTextWord(proDocPath)
        mainArray = txt.split('\n\n')  # Split by entered seperations
        for promptSegment in mainArray:  # Repeat for each block
            if promptSegment != "":  # If not blank
                if promptSegment[0:1] == "\n":  # If line starts with a line break
                    promptSegment = promptSegment[1:len(promptSegment)]  # Remove line break
                cutUp = promptSegment.split("\n")
                titleLen = len(cutUp[0])
                if titleLen < 25:
                    promptSegment = promptSegment[0:titleLen] + ":" + promptSegment[titleLen:]
                else:
                    txtDoc.write("Unnamed:\n")
                txtDoc.write(promptSegment)
                txtDoc.write("\n")
                txtDoc.write("\n")

        os.remove(proDocPath)

    txtDoc.close()



folderPathway = r"C:\Users\zaper\Downloads\OG PAGES DOCS\Found, G, VA"  # Declare main folder pathway
initializePrompt(folderPathway)  # Create/open the found prompt document
fileList = getDocx(folderPathway)  # Get all word documents from folder
for fileName in fileList:  # Repeat for each word doc
    filePath = folderPathway + chr(92) + fileName
    word2txtDoc(filePath)  # Convert word doc to a text doc


# Give final report
print(" ")
print("Remaining Files: ")
print(getDocx(folderPathway))
print('Complete')


